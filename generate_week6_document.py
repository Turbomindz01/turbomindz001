from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Turbomindz Week 6 Production', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('April 27 - May 3, 2026')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Spacing

# Week 6 Content
days_content = [
    {
        "day": "DAY 1: MONDAY",
        "date": "April 28, 2026",
        "title": "Printful Store Launch",
        "posting_time": "11 AM PT",
        "platforms": {
            "Instagram": """The merch isn't decoration. It's a statement.

We just launched the Printful store with 5 new items—mugs, shirts, and totes carrying our most philosophical scenes. TM-256 (Floating Sumi-e), TM-261 (Neon Calligraphy Void), and TM-270 (Clockwork Garden) are now wearable, holdable, shareable.

When you wear Turbomindz, you're wearing philosophy. You're saying: This art matters. These questions matter. This community matters.

Each product is printed on demand. No waste. No surplus. Just intentional creation, the way every scene in the 888-scene universe is intentional.

The philosophy starts on the blockchain. The belonging starts when you hold it.

Shop now: [Printful Store Link]

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#Printful #MerchLaunch #PhilosophyArt #NFTCommunity #DigitalArt #Wearable #ArtistLife #CreativeUniverse #Philosophy #BlockchainArt #PrintfulStore #ArtMerch #SupportArtists #IndieCreator #VisualPhilosophy #SceneArt #LimitedEdition #ArtCommunity #CreativeWorld #NFTArt #PhilosophicalArt #MerchStore #ArtSupport""",
            "Twitter/X": "Printful store live. Mugs, shirts, totes with your favorite scenes. TM-256, TM-261, TM-270 are now wearable philosophy. 🔮✨",
            "Discord": "🎨 PRINTFUL LAUNCH 🎨\nThe merch is here. Mugs, tees, and totes with scenes TM-256, TM-261, and TM-270. Each item printed on demand—no waste, pure intention. This is philosophy you can hold. Shop now and carry the universe with you.",
            "Pinterest Title": "Philosophical Art Merch - Sumi-e Inspired NFT Designs on Printful Mugs, Shirts & Totes",
            "Pinterest Description": "Turbomindz launches the Printful store with 5 new merch items featuring hand-crafted philosophical art scenes. Wearable NFT-inspired designs including Floating Sumi-e, Neon Calligraphy Void, and Clockwork Garden. Shop sustainable, on-demand printed apparel and accessories with deep artistic meaning."
        }
    },
    {
        "day": "DAY 2: TUESDAY",
        "date": "April 29, 2026",
        "title": "Scene Reveal TM-256",
        "posting_time": "10 AM PT",
        "platforms": {
            "Instagram": """Scene TM-256: Floating Sumi-e

Minimal. Ancient. Impossible.

Japanese ink wash painting translated into digital space where physics don't apply. Brushstrokes dissolve into emptiness. Emptiness blooms into meaning. This is what happens when centuries of Eastern aesthetic tradition meet the infinite canvas of possibility.

Sumi-e teaches us that what you leave out is as important as what you paint in. Silence is a color. Space is substance. And the artist's hand—even when painting digital—is an extension of breath and meditation.

TM-256 isn't trying to be beautiful. It's beautiful because it's honest. Because every stroke asks a question instead of declaring an answer.

The scene is live in the 888-scene universe. Own the poetry. Own the philosophy.

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#SceneReveal #TM256 #SumiE #JapaneseBrushwork #PhilosophyArt #DigitalArt #NFT #ArtisticPhilosophy #Minimalism #EasternAesthetics #BlockchainArt #CreativeUniverse #DigitalPainting #ArtisticExecution #Contemplative #ZenPhilosophy #ArtCommunity #SceneArt #PoetryInArt #VisualPhilosophy #NFTScene #Aesthetics #ArtSupport""",
            "Twitter/X": "Scene TM-256 drops: Floating Sumi-e. Japanese ink wash painting meets the infinite. Minimal. Honest. Pure. 🖌️✨",
            "Discord": "🎨 SCENE REVEAL: TM-256 🎨\nFloating Sumi-e is now live. Watch brushstrokes dissolve into emptiness. Watch emptiness bloom into meaning. This scene carries centuries of Eastern philosophy in every stroke. It asks: what matters when you remove everything else?",
            "Pinterest Title": "Floating Sumi-e - Digital NFT Art with Japanese Ink Wash Philosophy and Zen Aesthetics",
            "Pinterest Description": "Turbomindz Scene TM-256 combines traditional Japanese Sumi-e painting with digital NFT art. Minimal brushwork, contemplative design, and zen philosophy. Perfect for collectors seeking meaningful artistic NFTs with deep cultural roots and meditative beauty."
        }
    },
    {
        "day": "DAY 3: WEDNESDAY",
        "date": "April 30, 2026",
        "title": "Medium Article Promotion",
        "posting_time": "2 PM PT",
        "platforms": {
            "Instagram": """We wrote about philosophy. Inside every NFT you've ever seen.

Most people think NFTs are about ownership. About JPEG certificates. About having something digital.

But what if every NFT is actually a philosophical argument? What if just by creating or collecting digital art on the blockchain, you're making claims about authenticity, permanence, value, and meaning?

This week on Medium, we explore the hidden philosophy behind every NFT—the questions that philosophers like Plato, Benjamin, and Wittgenstein were wrestling with centuries ago. Questions that come alive when you tokenize art.

It's not about NFTs being good or bad. It's about understanding what they're actually saying about reality.

Read the full article: "The Hidden Philosophy Behind Every NFT You've Ever Seen"

[Link to Medium Article]

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#Philosophy #NFT #DigitalArt #Blockchain #Authenticity #ArtPhilosophy #MediumArticle #ThinkingAboutArt #CulturalCriticism #PleasureReading #ArtisticIntelligence #CreativeWriting #Essays #VisualPhilosophy #WhatIsValue #NFTCommunity #PhilosophyMatters #BlockchainPhilosophy #ArtAndThought #DeepDive""",
            "Twitter/X": "We wrote 1,500 words about the philosophy hiding inside every NFT. Plato, Benjamin, Wittgenstein. What does ownership really mean? Read on Medium 📖✨",
            "Discord": "📖 MEDIUM ARTICLE LIVE 📖\n\"The Hidden Philosophy Behind Every NFT You've Ever Seen\" is out now. We break down how every NFT carries philosophical assumptions about authenticity, permanence, ownership, and value. This isn't crypto talk—it's real philosophy. [Link]",
            "Pinterest Title": "Philosophy Behind NFTs - Blockchain Art & Digital Ownership Essays",
            "Pinterest Description": "Deep dive into the philosophy hidden inside NFT art and blockchain technology. Explore concepts of authenticity, permanence, and ownership through the lens of philosophers like Plato, Benjamin, and Wittgenstein. Thoughtful essays on digital art meaning."
        }
    },
    {
        "day": "DAY 4: THURSDAY",
        "date": "May 1, 2026",
        "title": "Behind-the-Scenes (Wallet Feature)",
        "posting_time": "9 AM PT",
        "platforms": {
            "Instagram": """How do you build trust in code?

This week we launched the thirdweb wallet integration on the Turbomindz app. Behind the scenes, it sounds simple: users connect their wallet, they authenticate, they own their scenes.

But the philosophy underneath? That's the real build.

Every time you hand someone the power to control their own cryptographic keys, you're trusting them with autonomy. You're saying: You own this. Not us. Not a platform. You.

The wallet isn't a feature. It's a statement. It's the technical manifestation of the biggest idea in Turbomindz: that philosophy lives in the choices we make, and art lives in the meaning we create together.

We spent weeks testing, iterating, rebuilding. Not because the code was hard. Because we wanted to get the values right.

This is what it looks like when technology serves philosophy instead of the other way around.

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#BehindTheScenes #Development #Blockchain #Wallet #Authentication #Technology #Philosophy #CreativeEngineering #Indiedev #StartupLife #TechForGood #Web3 #Crypto #UserOwnership #Autonomy #BuildingInPublic #DeveloperLife #CreativeCode #FutureOfWork #Innovation #TrustInCode""",
            "Twitter/X": "Behind the scenes: How do you build trust in code? This week we launched wallet auth on Turbomindz. It's philosophy in practice. 🔑✨",
            "Discord": "⚙️ BEHIND-THE-SCENES: WALLET LAUNCH ⚙️\nThirdweb wallet integration is now live on the app. Users authenticate, connect, and own their scenes. We spent weeks getting the values right—because technology should serve philosophy, not the other way around.",
            "Pinterest Title": "Web3 Wallet Authentication - Philosophy of Blockchain Ownership and User Autonomy",
            "Pinterest Description": "Behind-the-scenes look at integrating blockchain wallet authentication for digital art ownership. Thirdweb integration, user autonomy, cryptographic trust. Technical philosophy of decentralized art platforms and NFT authentication systems."
        }
    },
    {
        "day": "DAY 5: FRIDAY",
        "date": "May 2, 2026",
        "title": "Scene Reveal TM-261",
        "posting_time": "10 AM PT",
        "platforms": {
            "Instagram": """Scene TM-261: Neon Calligraphy Void

Characters of light. Cutting through silence.

Each stroke is a meditation. Eastern calligraphy—the art of making meaning visible through the body's movement—translated into neon against the void. The characters aren't just symbols; they're movements made permanent.

Calligraphy teaches that the hand is the mind. That there's no separation between thinking and moving, between intention and form. When you write with a brush—or now, with light—you're not just recording symbols. You're recording presence.

TM-261 asks: What happens when ancient wisdom meets digital void? When the shape of a character (years of tradition, breath, meditation) floats alone in impossible space?

The answer is luminescence. The answer is presence. The answer is understanding that emptiness isn't absence—it's potential.

Own the void. Own the light. Own the question.

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#SceneReveal #TM261 #Calligraphy #Neon #EasternWisdom #DigitalArt #NFT #LightArt #CharacterDesign #ArtisticPhilosophy #Meditation #Luminescence #VoidArt #CreativeUniverse #VisualPhilosophy #BlockchainArt #ArtCommunity #SceneArt #MeaningMaking #DigitalCalligraphy #NFTScene""",
            "Twitter/X": "Scene TM-261: Neon Calligraphy Void. Light carving through silence. Characters floating in potential. 🔮✨",
            "Discord": "🎨 SCENE REVEAL: TM-261 🎨\nNeon Calligraphy Void is live. Watch characters of light cut through emptiness. Each stroke carries Eastern wisdom. Each character asks: What happens when tradition meets the void? Own the luminescence.",
            "Pinterest Title": "Neon Calligraphy NFT Art - Digital Eastern Wisdom and Light-Based Digital Painting",
            "Pinterest Description": "Turbomindz Scene TM-261 combines traditional Eastern calligraphy with neon digital art. Glowing characters floating in void space. Contemplative light-based NFT design exploring meaning, presence, and philosophical depth through luminescent letterforms."
        }
    },
    {
        "day": "DAY 6: SATURDAY",
        "date": "May 3, 2026",
        "title": "Art Education (Sumi-e)",
        "posting_time": "12 PM PT",
        "platforms": {
            "Instagram": """What is Sumi-e? 🖌️

It's Japanese ink wash painting. But that's like saying the ocean is salt water. True, but missing everything that matters.

Sumi-e is a meditation practice. A spiritual discipline. A way of thinking about what's essential.

In Sumi-e, you:
- Use only black ink
- Paint on white paper
- Make each stroke count (you can't erase)
- Leave space as a color
- Let emptiness speak

The goal isn't photorealism. It's essence. The artist spends years learning to paint a bamboo shoot with three strokes—not because three strokes is all you have, but because three strokes contain everything.

This is why Scene TM-256 (Floating Sumi-e) matters. It brings this 800-year-old philosophy into the digital universe. It asks: In a world of infinite canvas, what do we still choose to leave empty?

Sumi-e teaches us that constraint breeds poetry. That simplicity is the hardest thing to make. That sometimes the most powerful statement is what you don't paint.

Want to learn more? Pick up a brush. Paint until you understand that less is infinite.

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#SumiE #JapanesePainting #ArtEducation #BrushworkMastery #ArtTechnique #PhilosophyOfArt #JapaneseCulture #ArtHistory #Minimalism #PaintingTechnique #ArtisticPhilosophy #EasternArt #Contemplative #ArtLessons #CulturalArt #Meditation #Aesthetics #VisualPhilosophy #ArtCommunity #LearnToArt""",
            "Twitter/X": "Sumi-e: Japanese ink painting where 3 strokes contain infinity. Constraint breeds poetry. Emptiness speaks. 🖌️✨",
            "Discord": "🎨 ART EDUCATION: SUMI-E 🎨\nWhat is Japanese ink wash painting? It's meditation in form. Constraint. Essence. Three strokes, infinite meaning. This is the philosophy behind Scene TM-256. Study the tradition. Understand the depth.",
            "Pinterest Title": "Sumi-E Japanese Ink Painting Tutorial - History, Technique, and Philosophical Art Practice",
            "Pinterest Description": "Learn about Sumi-E, the classical Japanese ink wash painting technique. Explore the philosophy of minimalism, brush mastery, and spiritual practice in traditional Eastern art. Guide to Sumi-E history, techniques, and meditative painting methods."
        }
    },
    {
        "day": "DAY 7: SUNDAY",
        "date": "May 4, 2026",
        "title": "Scene Reveal TM-270 + Weekly Recap",
        "posting_time": "6 PM PT",
        "platforms": {
            "Instagram": """Scene TM-270: Clockwork Garden 🕰️🌿

Gears and blooms. Time and growth. The machine dreaming in flowers.

This scene closes Week 6 with the hardest question: What happens when efficiency meets beauty? When the precision of clockwork collides with the chaos of gardens?

In the industrial world, we learned to separate these. Machines are for output. Gardens are for contemplation. Work is different from rest. Productivity is the opposite of presence.

But TM-270 whispers something different. It says: What if the machine becomes the garden? What if gears can grow? What if time can bloom?

This is the philosophical core of Mega-Prompt 3 (Wallet & Auth): We've built tools. Technical infrastructure. Code that authenticates. But the real work is making technology serve meaning, making systems cultivate beauty, making efficiency hold space for mystery.

The scenes in Turbomindz aren't escaping technology. They're redeeming it.

WEEK 6 RECAP:
- Wallet + Auth launched (you own your scenes now)
- Printful merch goes live (philosophy becomes wearable)
- Medium article published (philosophy in 1,500 words)
- Scenes TM-256, TM-261, TM-270 featured (poetry in pixels)
- Community asking better questions

Week 7 is coming. The next Mega-Prompt. The next impossible question.

Keep the garden. Keep the gears. Keep looking up.

━━━━━━━━━━━━━━━━━━━━━━━━━━
💫 The NFT is the soul
🎨 The art is the body
🏘 The community is the village
📖 The story is the mirror
━━━━━━━━━━━━━━━━━━━━━━━━━━

#SceneReveal #TM270 #ClockworkGarden #PhilosophyArt #Enlightenment #TimelessDesign #GearArt #NatureAndTechnology #DigitalArt #NFT #CreativeUniverse #MechanicalBeauty #SteampunkPhilosophy #ArtisticVision #WeeklyRecap #CommunityLove #ArtSupport #BlockchainArt #VisualPhilosophy #DreamingMachines""",
            "Twitter/X": "Scene TM-270: Clockwork Garden. Gears blooming. Time flowering. What if machines dream in beauty? 🕰️✨",
            "Discord": "🎨 SCENE REVEAL: TM-270 🎨\nClockwork Garden closes Week 6. Gears and blooms. Machine and growth. The hardest question: What happens when efficiency meets beauty? Own the dream. Join the garden. Week 7 is coming.\n\nThanks for being here this week. Keep the philosophy. Keep the art. Keep looking up.",
            "Pinterest Title": "Clockwork Garden - Steampunk NFT Art with Enlightenment Philosophy and Mechanical Beauty",
            "Pinterest Description": "Turbomindz Scene TM-270 explores the intersection of technology and nature through steampunk-inspired NFT art. Gears, blooms, and enlightenment philosophy. Meditative design merging mechanical precision with organic beauty and spiritual contemplation."
        }
    }
]

# Add content for each day
for day_num, day in enumerate(days_content, 1):
    # Day header
    heading = doc.add_heading(f"{day['day']} – {day['title']}", level=1)
    
    # Date and posting time
    info_para = doc.add_paragraph()
    info_para.add_run(f"📅 Date: ").font.bold = True
    info_para.add_run(day['date'])
    
    info_para2 = doc.add_paragraph()
    info_para2.add_run(f"🕐 Posting Time: ").font.bold = True
    info_para2.add_run(day['posting_time'])
    
    doc.add_paragraph()  # Spacing
    
    # Platform-specific content
    for platform, content in day['platforms'].items():
        # Platform header
        platform_heading = doc.add_heading(platform, level=2)
        platform_heading_format = platform_heading.runs[0]
        platform_heading_format.font.color.rgb = RGBColor(0, 102, 153)
        
        # Content
        content_para = doc.add_paragraph(content)
        content_para.paragraph_format.line_spacing = 1.15
        
        doc.add_paragraph()  # Spacing
    
    # Add page break between days (except the last one)
    if day_num < len(days_content):
        doc.add_page_break()

# Save the document
output_path = r'c:\Users\hugog\Desktop\TMDZ\Turbomindz_Week6_Production.docx'
doc.save(output_path)
print(f"✅ Document created successfully: {output_path}")
print(f"📄 Total pages: ~{len(days_content)}")
print(f"📊 Content includes:")
print(f"   • 7 days of posts (Monday-Sunday)")
print(f"   • Instagram captions with hashtags")
print(f"   • Twitter/X posts")
print(f"   • Discord announcements")
print(f"   • Pinterest titles and descriptions")
print(f"   • Posting times and dates")
